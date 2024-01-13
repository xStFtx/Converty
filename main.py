import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading
import ebooklib
from ebooklib import epub
import pdfkit
from pdf2docx import parse
import os
import logging
from docx import Document

# Setup logging
logging.basicConfig(filename='conversion.log', level=logging.INFO, format='%(asctime)s:%(levelname)s:%(message)s')

# Ensure you have the correct path to the wkhtmltopdf executable
wkhtmltopdf_path = r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe'
config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)

# Progress Bar Update Function
def update_progress_bar(progress_bar, value):
    progress_bar['value'] = value
    root.update_idletasks()

# Convert EPUB to PDF Function with threading
def threaded_convert_epub_to_pdf(file_path, save_path, progress_bar):
    try:
        epub_book = epub.read_epub(file_path)
        html_content_list = []

        for item in epub_book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                html_content_list.append(item.get_content())

        temp_file_path = 'temp.html'
        with open(temp_file_path, 'wb') as html_file:
            for html_content in html_content_list:
                html_file.write(html_content)

        pdfkit.from_file(temp_file_path, save_path, configuration=config)
        logging.info(f"Converted {file_path} to {save_path}")
        messagebox.showinfo("Success", f"File saved as {save_path}")
    except Exception as e:
        logging.error(f"Error converting {file_path}: {e}")
        messagebox.showerror("Conversion Error", str(e))
    finally:
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)
        update_progress_bar(progress_bar, 0)

# Convert PDF to DOCX Function with threading
def threaded_convert_pdf_to_docx(file_path, save_path, progress_bar):
    try:
        parse(file_path, save_path)
        logging.info(f"Converted {file_path} to {save_path}")
        messagebox.showinfo("Success", f"File saved as {save_path}")
    except Exception as e:
        logging.error(f"Error converting {file_path}: {e}")
        messagebox.showerror("Conversion Error", str(e))
    finally:
        update_progress_bar(progress_bar, 0)

# Convert TXT to DOCX Function with threading
def threaded_convert_txt_to_docx(file_path, save_path, progress_bar):
    try:
        document = Document()
        with open(file_path, 'r', encoding='utf-8') as file:
            for line in file:
                document.add_paragraph(line)
        document.save(save_path)
        logging.info(f"Converted {file_path} to {save_path}")
        messagebox.showinfo("Success", f"File saved as {save_path}")
    except Exception as e:
        logging.error(f"Error converting {file_path}: {e}")
        messagebox.showerror("Conversion Error", str(e))
    finally:
        update_progress_bar(progress_bar, 0)

# Convert DOCX to TXT Function with threading
def threaded_convert_docx_to_txt(file_path, save_path, progress_bar):
    try:
        document = Document(file_path)
        with open(save_path, 'w', encoding='utf-8') as file:
            for para in document.paragraphs:
                file.write(para.text + '\n')
        logging.info(f"Converted {file_path} to {save_path}")
        messagebox.showinfo("Success", f"File saved as {save_path}")
    except Exception as e:
        logging.error(f"Error converting {file_path}: {e}")
        messagebox.showerror("Conversion Error", str(e))
    finally:
        update_progress_bar(progress_bar, 0)

# Wrapper functions for Threading
def convert_epub_to_pdf_wrapper(progress_bar):
    file_path = filedialog.askopenfilename(filetypes=[("EPUB files", "*.epub")])
    if file_path:
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        save_path = os.path.join(os.path.dirname(file_path), f"{base_name}.pdf")
        update_progress_bar(progress_bar, 20)
        threading.Thread(target=
    threaded_convert_epub_to_pdf, args=(file_path, save_path, progress_bar)).start()

def convert_pdf_to_docx_wrapper(progress_bar):
    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
    if file_path:
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        save_path = os.path.join(os.path.dirname(file_path), f"{base_name}.docx")
        update_progress_bar(progress_bar, 20)
        threading.Thread(target=threaded_convert_pdf_to_docx, args=(file_path, save_path, progress_bar)).start()

def convert_txt_to_docx_wrapper(progress_bar):
    file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt")])
    if file_path:
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        save_path = os.path.join(os.path.dirname(file_path), f"{base_name}.docx")
        update_progress_bar(progress_bar, 20)
        threading.Thread(target=threaded_convert_txt_to_docx, args=(file_path, save_path, progress_bar)).start()

def convert_docx_to_txt_wrapper(progress_bar):
    file_path = filedialog.askopenfilename(filetypes=[("Word documents", "*.docx")])
    if file_path:
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        save_path = os.path.join(os.path.dirname(file_path), f"{base_name}.txt")
        update_progress_bar(progress_bar, 20)
        threading.Thread(target=threaded_convert_docx_to_txt, args=(file_path, save_path, progress_bar)).start()

# GUI Setup
root = tk.Tk()
root.title("File Format Converter")

frame = tk.Frame(root)
frame.pack(padx=10, pady=10)

progress = ttk.Progressbar(frame, orient=tk.HORIZONTAL, length=200, mode='determinate')
progress.pack(pady=(0, 10))

epub_to_pdf_button = tk.Button(frame, text="Convert EPUB to PDF", command=lambda: convert_epub_to_pdf_wrapper(progress))
epub_to_pdf_button.pack(pady=(0, 5))

pdf_to_docx_button = tk.Button(frame, text="Convert PDF to DOCX", command=lambda: convert_pdf_to_docx_wrapper(progress))
pdf_to_docx_button.pack(pady=(0, 5))

txt_to_docx_button = tk.Button(frame, text="Convert TXT to DOCX", command=lambda: convert_txt_to_docx_wrapper(progress))
txt_to_docx_button.pack(pady=(0, 5))

docx_to_txt_button = tk.Button(frame, text="Convert DOCX to TXT", command=lambda: convert_docx_to_txt_wrapper(progress))
docx_to_txt_button.pack()

root.mainloop()
